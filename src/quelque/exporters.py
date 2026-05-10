from __future__ import annotations

import io

from docx import Document
from fpdf import FPDF

from .schemas import LectureNotesResult


def render_markdown(result: LectureNotesResult) -> str:
    lines = [
        "# Quelque Notes",
        "",
        "## Summary",
        result.summary,
        "",
        "## Key Takeaways",
    ]
    lines.extend(f"- {item}" for item in result.key_takeaways)
    lines.extend(["", "## Study Notes"])
    lines.extend(f"- {item}" for item in result.study_notes)
    lines.extend(["", "## Glossary"])
    lines.extend(f"- {item}" for item in result.glossary)
    lines.extend(["", "## Action Items"])
    lines.extend(f"- {item}" for item in result.action_items)
    lines.extend(["", "## Sanitized Transcript", result.sanitized_transcript])
    return "\n".join(lines)


def render_docx(result: LectureNotesResult) -> bytes:
    doc = Document()
    doc.add_heading("Quelque Notes", level=1)
    _add_section(doc, "Summary", [result.summary])
    _add_section(doc, "Key Takeaways", result.key_takeaways)
    _add_section(doc, "Study Notes", result.study_notes)
    _add_section(doc, "Glossary", result.glossary)
    _add_section(doc, "Action Items", result.action_items)
    _add_section(doc, "Sanitized Transcript", [result.sanitized_transcript])
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def render_pdf(result: LectureNotesResult) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    _add_pdf_section(pdf, "Summary", [result.summary])
    _add_pdf_section(pdf, "Key Takeaways", result.key_takeaways)
    _add_pdf_section(pdf, "Study Notes", result.study_notes)
    _add_pdf_section(pdf, "Glossary", result.glossary)
    _add_pdf_section(pdf, "Action Items", result.action_items)
    _add_pdf_section(pdf, "Sanitized Transcript", [result.sanitized_transcript])
    return bytes(pdf.output())


def _add_section(doc: Document, title: str, items: list[str]) -> None:
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(item)


def _add_pdf_section(pdf: FPDF, title: str, items: list[str]) -> None:
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=12)
    for item in items:
        # Ensure compatibility with standard Helvetica latin-1 and avoid width bugs
        safe_item = item.encode("latin-1", "replace").decode("latin-1")
        # Manually break excessively long unbroken strings (URLs, dashes) that crash FPDF
        words = []
        for word in safe_item.split(" "):
            if len(word) > 60:
                words.append(" ".join(word[i:i+60] for i in range(0, len(word), 60)))
            else:
                words.append(word)
        safe_item = " ".join(words)
        try:
            pdf.multi_cell(0, 8, safe_item)
        except Exception:
            pass
    pdf.ln(2)
